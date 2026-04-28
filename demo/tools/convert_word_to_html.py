"""
将Word文档转换为HTML格式的工具脚本
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import html

def parse_docx_to_html(docx_path):
    """将Word文档解析为HTML"""
    doc = Document(docx_path)
    html_content = []

    html_content.append('<!DOCTYPE html>')
    html_content.append('<html lang="zh-CN">')
    html_content.append('<head>')
    html_content.append('    <meta charset="utf-8">')
    html_content.append(f'    <title>{os.path.basename(docx_path)}</title>')
    html_content.append('    <style>')
    html_content.append('        body { font-family: "Microsoft YaHei", Arial, sans-serif; line-height: 1.8; color: #333; padding: 20px; }')
    html_content.append('        h1 { font-size: 24px; font-weight: bold; color: #2c3e50; margin: 30px 0 20px 0; border-bottom: 3px solid #3498db; }')
    html_content.append('        h2 { font-size: 20px; font-weight: bold; color: #34495e; margin: 25px 0 15px 0; border-left: 4px solid #3498db; padding-left: 10px; }')
    html_content.append('        h3 { font-size: 18px; font-weight: bold; color: #555; margin: 20px 0 10px 0; }')
    html_content.append('        p { margin: 10px 0; }')
    html_content.append('        table { width: 100%; border-collapse: collapse; margin: 20px 0; }')
    html_content.append('        table, th, td { border: 1px solid #444; }')
    html_content.append('        th { background-color: #3498db; color: white; padding: 12px 8px; text-align: left; }')
    html_content.append('        td { padding: 10px 8px; vertical-align: top; }')
    html_content.append('        tr:nth-child(even) { background-color: #f8f9fa; }')
    html_content.append('        ul, ol { margin: 10px 0 10px 30px; }')
    html_content.append('        li { margin: 5px 0; }')
    html_content.append('        strong { font-weight: bold; color: #2c3e50; }')
    html_content.append('    </style>')
    html_content.append('</head>')
    html_content.append('<body>')

    # 解析段落和表格
    in_table = False
    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            # 处理表格
            table = next((t for t in doc.tables if t._element == element), None)
            if table:
                html_content.append('    <table>')
                for i, row in enumerate(table.rows):
                    html_content.append('        <tr>')
                    for cell in row.cells:
                        text = cell.text.strip()
                        # 第一行通常是表头
                        tag = 'th' if i == 0 else 'td'
                        html_content.append(f'            <{tag}>{html.escape(text)}</{tag}>')
                    html_content.append('        </tr>')
                html_content.append('    </table>')
        elif element.tag.endswith('p'):
            # 处理段落
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if para:
                text = para.text.strip()
                if not text:
                    continue

                # 根据样式判断标题级别
                style_name = para.style.name.lower()
                is_bold = any(run.bold for run in para.runs if run.text.strip())

                if 'heading 1' in style_name or (is_bold and len(text) < 30 and text.startswith('一、')):
                    html_content.append(f'    <h1>{html.escape(text)}</h1>')
                elif 'heading 2' in style_name or (is_bold and len(text) < 30 and text.startswith(('二、', '三、', '四、', '五、'))):
                    html_content.append(f'    <h2>{html.escape(text)}</h2>')
                elif 'heading 3' in style_name or (is_bold and len(text) < 40):
                    html_content.append(f'    <h3>{html.escape(text)}</h3>')
                elif is_bold:
                    html_content.append(f'    <p><strong>{html.escape(text)}</strong></p>')
                else:
                    # 检查是否是列表项
                    if text.startswith(('- ', '• ', '1.', '2.', '3.', '①', '②', '③')):
                        html_content.append(f'    <li>{html.escape(text[2:].strip() if text[1] in (" ", ".") else text)}</li>')
                    else:
                        html_content.append(f'    <p>{html.escape(text)}</p>')

    html_content.append('</body>')
    html_content.append('</html>')

    return '\n'.join(html_content)


def convert_all_word_docs():
    """转换所有Word文档"""
    template_dir = r'D:\1work\20251016\breast_health_mvp\模板'
    output_dir = r'D:\1work\20251016\breast_health_mvp\tools\converted_html'

    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)

    # 需要转换的文档
    word_files = [
        '肺结节健康管理方案模板.docx',
        '甲状腺结节健康管理方案模板.docx',
        '乳腺合并肺双结节健康管理方案模板.docx',
        '乳腺合并甲状腺双结节健康管理方案模板.docx',
        '甲状腺合并肺双结节健康管理方案模板.docx',
        '三结节健康管理方案方案模板.docx'
    ]

    for word_file in word_files:
        docx_path = os.path.join(template_dir, word_file)
        if os.path.exists(docx_path):
            print(f'Converting: {word_file}')
            html_content = parse_docx_to_html(docx_path)

            # 保存HTML
            output_file = word_file.replace('.docx', '.html')
            output_path = os.path.join(output_dir, output_file)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            print(f'Saved: {output_path}')
        else:
            print(f'File not found: {docx_path}')


if __name__ == '__main__':
    try:
        convert_all_word_docs()
        print('\nAll Word documents converted successfully!')
    except Exception as e:
        print(f'Conversion failed: {e}')
        import traceback
        traceback.print_exc()
